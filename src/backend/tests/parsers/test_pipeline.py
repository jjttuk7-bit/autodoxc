"""DA4 파서 디스패처 + docx + pdf_text 통합 테스트."""
from __future__ import annotations

from pathlib import Path

from app.parsers import parse
from app.parsers.pipeline import detect_format


def test_detect_format_known(tmp_path: Path) -> None:
    assert detect_format("a.docx") == "docx"
    assert detect_format("a.pdf") == "pdf_text"
    assert detect_format("a.hwpx") == "hwpx"
    assert detect_format("a.hwp") == "hwp"
    assert detect_format("a.txt") == "txt"
    assert detect_format("a.bin") == "unknown"


def test_txt_passthrough(tmp_path: Path) -> None:
    p = tmp_path / "memo.txt"
    p.write_text("간단한 메모", encoding="utf-8")
    result = parse(p)
    assert result.format == "txt"
    assert result.raw_text == "간단한 메모"
    assert len(result.blocks) == 1


def test_unsupported_returns_error(tmp_path: Path) -> None:
    p = tmp_path / "binary.bin"
    p.write_bytes(b"\x00\x01\x02")
    result = parse(p)
    assert result.format == "unknown"
    assert result.confidence == 0.0
    assert any(w.severity == "error" for w in result.warnings)


def test_docx_real_file(tmp_path: Path) -> None:
    """python-docx로 문서 생성 → 파싱 라운드트립."""
    from docx import Document

    p = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("1. 고용 대상 업종", level=1)
    doc.add_paragraph("본 사는 항공우주 부품 제조업체이며,")
    doc.add_paragraph("5축 가공 분야의 특수 합금 가공이 핵심 역량입니다.")
    doc.save(str(p))

    result = parse(p)
    assert result.format == "docx"
    assert result.confidence >= 0.9
    assert any(b.kind == "heading" and "고용 대상" in b.text for b in result.blocks)
    assert "항공우주" in result.raw_text
