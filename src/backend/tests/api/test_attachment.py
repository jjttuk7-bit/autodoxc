"""첨부 양식 업로드 → 골격 추출 엔드포인트."""
from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient

from app.api.sessions import _test_clear_store, _test_get_context
from app.main import app


@pytest.fixture(autouse=True)
def _clear_store():
    _test_clear_store()
    yield
    _test_clear_store()


def _new_session(client: TestClient) -> str:
    resp = client.post("/api/sessions", json={"user_input": "양식 첨부 테스트"})
    return resp.json()["session_id"]


def _docx_bytes(headings: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for h in headings:
        doc.add_heading(h, level=1)
        doc.add_paragraph(f"{h}에 대한 안내 문단.")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def test_upload_docx_extracts_skeleton() -> None:
    client = TestClient(app)
    sid = _new_session(client)
    content = _docx_bytes(["신청인", "신청 사유", "첨부 서류"])

    resp = client.post(
        f"/api/sessions/{sid}/attachment",
        files={
            "file": (
                "form.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert resp.status_code == 201
    body = resp.json()
    assert body["extracted"] is True
    assert len(body["section_titles"]) == 3
    assert "신청인" in body["section_titles"][0]
    # 세션에 attached_skeleton 저장됐는지
    ctx = _test_get_context(sid)
    assert ctx is not None
    assert len(ctx.attached_skeleton) == 3
    assert ctx.attached_skeleton[0].source.kind == "user_attached"


def test_upload_rejects_unsupported_format() -> None:
    client = TestClient(app)
    sid = _new_session(client)
    resp = client.post(
        f"/api/sessions/{sid}/attachment",
        files={"file": ("image.png", b"\x89PNG\r\n", "image/png")},
    )
    assert resp.status_code == 415


def test_upload_txt_without_headings_reports_not_extracted() -> None:
    client = TestClient(app)
    sid = _new_session(client)
    resp = client.post(
        f"/api/sessions/{sid}/attachment",
        files={"file": ("note.txt", "제목 없는 평문".encode("utf-8"), "text/plain")},
    )
    assert resp.status_code == 201
    body = resp.json()
    assert body["extracted"] is False
    assert body["section_titles"] == []
    assert body["warnings"]  # 폴백 안내 메시지


def test_upload_404_when_session_missing() -> None:
    client = TestClient(app)
    resp = client.post(
        "/api/sessions/nope/attachment",
        files={"file": ("a.txt", b"x", "text/plain")},
    )
    assert resp.status_code == 404
