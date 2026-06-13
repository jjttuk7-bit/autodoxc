"""세션 생성·스트리밍·편집 — B1-1, B1-4.

흐름:
1. POST /api/sessions  { user_input }                  → { session_id }
2. GET  /api/sessions/{id}/stream                      → SSE (오케스트레이터)
3. POST /api/sessions/{id}/fill   { section_id, paragraph_idx, text }
4. POST /api/sessions/{id}/answer { field_ids[], value, skip? }

세션 영속: 메모리 dict의 SessionContext (DB 없이 동작).
오케스트레이터가 결과를 SessionContext에 축적 → 후속 fill/answer가 그 위에 작용.
"""
from __future__ import annotations

import io
import re
import uuid
from collections.abc import AsyncGenerator
from dataclasses import dataclass, field
from datetime import datetime, timezone

from fastapi import APIRouter, HTTPException, Response, status
from pydantic import BaseModel, Field
from sse_starlette.sse import EventSourceResponse

from app.agents import DraftWriter
from app.llm import get_llm_client
from app.orchestrator import run_main_sequence
from app.shared.types import (
    DocType,
    DraftSection,
    Fact,
    ParagraphAnnotation,
    SkeletonNode,
)
from app.shared.types.agents.draft_writer import DraftWriterInput


router = APIRouter(prefix="/api/sessions", tags=["sessions"])


# --- 메모리 세션 store -----------------------------------------------------


@dataclass
class SessionContext:
    """세션의 누적 상태. 오케스트레이터·편집 endpoint가 공유."""
    session_id: str
    user_input: str
    doc_type: DocType | None = None
    skeleton: list[SkeletonNode] = field(default_factory=list)
    facts: list[Fact] = field(default_factory=list)
    sections: dict[str, DraftSection] = field(default_factory=dict)
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))

    def touch(self) -> None:
        self.updated_at = datetime.now(timezone.utc)


_SESSIONS: dict[str, SessionContext] = {}


def _new_session_id() -> str:
    return uuid.uuid4().hex[:16]


def _require_session(session_id: str) -> SessionContext:
    ctx = _SESSIONS.get(session_id)
    if ctx is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"session {session_id} not found",
        )
    return ctx


# --- 모델 ----------------------------------------------------------------


class CreateSessionRequest(BaseModel):
    user_input: str = Field(..., min_length=1, max_length=10000)


class CreateSessionResponse(BaseModel):
    session_id: str


class FillRequest(BaseModel):
    section_id: str
    paragraph_idx: int = Field(ge=0)
    text: str = Field(..., min_length=1, max_length=5000)


class FillResponse(BaseModel):
    section: DraftSection


class AnswerRequest(BaseModel):
    field_ids: list[str] = Field(default_factory=list)
    value: str = ""
    skip: bool = False


class AnswerResponse(BaseModel):
    acknowledged: bool
    facts_added: int
    skipped: bool
    updated_sections: list[DraftSection] = Field(default_factory=list)


class SessionStateResponse(BaseModel):
    session_id: str
    user_input: str
    doc_type: DocType | None
    skeleton: list[SkeletonNode]
    facts: list[Fact]
    sections: list[DraftSection]
    is_complete: bool
    updated_at: datetime


# --- POST /api/sessions ---------------------------------------------------


@router.post(
    "",
    response_model=CreateSessionResponse,
    status_code=status.HTTP_201_CREATED,
)
async def create_session(body: CreateSessionRequest) -> CreateSessionResponse:
    sid = _new_session_id()
    _SESSIONS[sid] = SessionContext(session_id=sid, user_input=body.user_input)
    return CreateSessionResponse(session_id=sid)


# --- GET /api/sessions/{id}/stream ---------------------------------------


@router.get("/{session_id}/stream")
async def stream_session(session_id: str):
    ctx = _require_session(session_id)
    headers = {
        "Cache-Control": "no-cache, no-transform",
        "X-Accel-Buffering": "no",
        "Connection": "keep-alive",
    }
    return EventSourceResponse(_orchestrate(ctx), headers=headers)


async def _orchestrate(ctx: SessionContext) -> AsyncGenerator[dict, None]:
    async for event in run_main_sequence(
        session_id=ctx.session_id,
        user_input=ctx.user_input,
        on_skeleton=lambda dt, sk: _set_skeleton(ctx, dt, sk),
        on_facts=lambda fs: _set_facts(ctx, fs),
        on_section=lambda s: _upsert_section(ctx, s),
    ):
        yield {"event": "message", "data": event.model_dump_json()}


def _set_skeleton(
    ctx: SessionContext, doc_type: DocType, skeleton: list[SkeletonNode]
) -> None:
    ctx.doc_type = doc_type
    ctx.skeleton = skeleton
    ctx.touch()


def _set_facts(ctx: SessionContext, facts: list[Fact]) -> None:
    ctx.facts = list(facts)
    ctx.touch()


def _upsert_section(ctx: SessionContext, section: DraftSection) -> None:
    ctx.sections[section.skeleton_id] = section
    ctx.touch()


# --- POST /api/sessions/{id}/fill -----------------------------------------


@router.post("/{session_id}/fill", response_model=FillResponse)
async def fill_slot(session_id: str, body: FillRequest) -> FillResponse:
    ctx = _require_session(session_id)
    section = ctx.sections.get(body.section_id)
    if section is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"section {body.section_id} not found",
        )
    paragraphs = list(section.paragraphs)
    if body.paragraph_idx >= len(paragraphs):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"paragraph_idx {body.paragraph_idx} out of range",
        )
    target = paragraphs[body.paragraph_idx]
    new_paragraph = target.model_copy(
        update={
            "text": body.text.strip(),
            "annotations": target.annotations.model_copy(
                update={
                    "status": "confirmed",
                    "needs_user_input": False,
                }
            ),
        }
    )
    paragraphs[body.paragraph_idx] = new_paragraph
    new_section = section.model_copy(update={"paragraphs": paragraphs})
    ctx.sections[body.section_id] = new_section
    ctx.touch()
    return FillResponse(section=new_section)


# --- POST /api/sessions/{id}/answer ---------------------------------------


@router.post("/{session_id}/answer", response_model=AnswerResponse)
async def answer_question(
    session_id: str, body: AnswerRequest
) -> AnswerResponse:
    ctx = _require_session(session_id)
    if body.skip:
        ctx.touch()
        return AnswerResponse(acknowledged=True, facts_added=0, skipped=True)

    # 1) facts에 추가
    added = 0
    for fid in body.field_ids:
        ctx.facts.append(
            Fact(
                field_id=fid,
                value=body.value,
                source="explicit",
                confidence=1.0,
            )
        )
        added += 1

    # 2) 영향 섹션 식별 → LLM으로 부분 재작성
    target_section_ids = _affected_sections(ctx, body.field_ids)
    updated_sections: list[DraftSection] = []
    if target_section_ids and ctx.doc_type and ctx.skeleton:
        writer = DraftWriter(get_llm_client())
        write_input = DraftWriterInput(
            skeleton=ctx.skeleton,
            facts=ctx.facts,
            doc_type=ctx.doc_type,
            target_sections=target_section_ids,
            force_llm=True,  # 시드 doc_type도 LLM으로 재작성 (facts 반영)
        )
        try:
            out = await writer.run(write_input)
            for section in out.draft.sections:
                ctx.sections[section.skeleton_id] = section
                updated_sections.append(section)
        except Exception:
            pass  # 재작성 실패 시 facts만 추가하고 응답

    ctx.touch()
    return AnswerResponse(
        acknowledged=True,
        facts_added=added,
        skipped=False,
        updated_sections=updated_sections,
    )


# --- 부분 재작성 영향 섹션 매핑 ------------------------------------------


_FIELD_TO_SECTION: dict[str, dict[str, list[str]]] = {
    "foreign-worker-employment-plan": {
        "recruitment_cost": ["sec_2"],
        "industry": ["sec_1"],
        "core_skill": ["sec_2"],
    },
    "administrative-appeal": {
        "disposition_date": ["sec_2", "sec_3"],
        "claimant_name": ["sec_1"],
        "respondent_name": ["sec_1"],
    },
    "content-certified-mail": {
        "recipient_name": ["sec_1", "sec_4"],
        "sender_name": ["sec_1"],
        "contract_date": ["sec_2"],
    },
}


def _affected_sections(ctx: SessionContext, field_ids: list[str]) -> list[str]:
    """답변된 field_id가 영향을 주는 섹션 id 목록.

    매핑된 field가 있으면 그 섹션만, 없으면 영향 0 (재작성 안 함).
    """
    if not ctx.doc_type or not ctx.skeleton:
        return []
    doc_map = _FIELD_TO_SECTION.get(ctx.doc_type.id, {})
    affected: set[str] = set()
    for fid in field_ids:
        affected.update(doc_map.get(fid, []))
    if not affected:
        # 매핑 없을 때: 최대 2개 섹션만 재작성 (비용·시간 균형)
        # — 미지의 field는 가장 빈 슬롯 많은 섹션 우선
        empty_counts: list[tuple[str, int]] = []
        for node in ctx.skeleton:
            section = ctx.sections.get(node.id)
            if section is None:
                continue
            empties = sum(
                1 for p in section.paragraphs if p.annotations.status == "empty"
            )
            if empties > 0:
                empty_counts.append((node.id, empties))
        empty_counts.sort(key=lambda x: -x[1])
        return [sid for sid, _ in empty_counts[:2]]
    return [sid for sid in affected if any(n.id == sid for n in ctx.skeleton)]


# --- GET /api/sessions/{id}/state -----------------------------------------


@router.get("/{session_id}/state", response_model=SessionStateResponse)
async def get_session_state(session_id: str) -> SessionStateResponse:
    """페이지 새로고침 시 복구용. SessionContext 전체 스냅샷."""
    ctx = _require_session(session_id)
    # 섹션을 skeleton 순서대로 정렬해서 반환
    ordered_sections: list[DraftSection] = []
    for node in ctx.skeleton:
        s = ctx.sections.get(node.id)
        if s is not None:
            ordered_sections.append(s)
    return SessionStateResponse(
        session_id=ctx.session_id,
        user_input=ctx.user_input,
        doc_type=ctx.doc_type,
        skeleton=ctx.skeleton,
        facts=ctx.facts,
        sections=ordered_sections,
        is_complete=len(ordered_sections) == len(ctx.skeleton)
        and len(ctx.skeleton) > 0,
        updated_at=ctx.updated_at,
    )


# --- GET /api/sessions/{id}/export?format=docx ----------------------------


@router.get("/{session_id}/export")
async def export_session(session_id: str, format: str = "docx") -> Response:
    """완성된 문서를 .docx로 다운로드."""
    ctx = _require_session(session_id)
    if format != "docx":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"unsupported format: {format}",
        )
    if not ctx.skeleton or not ctx.sections:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="document not ready for export — skeleton or sections missing",
        )
    bio = _build_docx(ctx)
    title = (ctx.doc_type.ko_name if ctx.doc_type else "행정문서").replace(
        " ", "_"
    )
    filename = f"autodoxc-{title}-{session_id}.docx"
    return Response(
        content=bio.getvalue(),
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            # 한글 파일명 RFC 5987
            "Content-Disposition": (
                f"attachment; filename=\"autodoxc-{session_id}.docx\"; "
                f"filename*=UTF-8''{filename}"
            )
        },
    )


def _build_docx(ctx: SessionContext) -> io.BytesIO:
    """python-docx로 SessionContext → 워드 문서 생성."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # 제목
    title_text = ctx.doc_type.ko_name if ctx.doc_type else "행정문서"
    title = doc.add_heading(title_text, level=0)
    for run in title.runs:
        run.font.size = Pt(20)

    # 사용자 입력 메타 (작은 글씨)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"생성: {ctx.updated_at.strftime('%Y-%m-%d %H:%M')} · 세션 {ctx.session_id}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)

    # 본문 — skeleton 순서대로
    for node in ctx.skeleton:
        section = ctx.sections.get(node.id)
        if section is None:
            continue
        # 섹션 제목
        doc.add_heading(section.title, level=1)
        # 문단들 — 자리표시자 [[필드명]]은 (필드명)으로 변환
        for p in section.paragraphs:
            para = doc.add_paragraph()
            converted = _convert_placeholders_for_docx(p.text)
            run = para.add_run(converted)
            if p.annotations.status == "inferred":
                run.italic = True
            elif p.annotations.status == "defaulted":
                run.font.size = Pt(10)
            elif p.annotations.status == "empty":
                run.italic = True
                run.bold = False

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


_PLACEHOLDER_REGEX = re.compile(r"\[\[([^\]]+)\]\]")


def _convert_placeholders_for_docx(text: str) -> str:
    """[[필드명]] → ( 필드명          ) 행정문서 표준 빈 칸 표기.

    출력본을 인쇄해 수기로 채울 수 있게 괄호 + 밑줄 형태.
    """
    return _PLACEHOLDER_REGEX.sub(lambda m: f"(  {m.group(1)}          )", text)


# --- 테스트·디버깅 유틸 ---


def _test_clear_store() -> None:
    _SESSIONS.clear()


def _test_get_context(session_id: str) -> SessionContext | None:
    return _SESSIONS.get(session_id)


# 구버전 호환 alias — 기존 테스트가 _test_get_input을 호출
def _test_get_input(session_id: str) -> str | None:
    ctx = _SESSIONS.get(session_id)
    return ctx.user_input if ctx else None


# Annotation 사용 명시 (lint 회피, 향후 활용)
_ = ParagraphAnnotation
