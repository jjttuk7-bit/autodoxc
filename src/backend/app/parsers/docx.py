"""DOCX 파서 — python-docx 기반.

ADR 0006 우선순위 1순위 포맷. 가장 안정적.
"""
from __future__ import annotations

from pathlib import Path

from .pipeline import ParsedBlock, ParseResult, ParseWarning


def parse_docx(path: Path, *, attachment_id: str) -> ParseResult:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        return ParseResult(
            attachment_id=attachment_id,
            format="docx",
            raw_text="",
            warnings=[
                ParseWarning(severity="error", message=f"python-docx 미설치: {e!s}")
            ],
            confidence=0.0,
        )

    try:
        doc = Document(str(path))
    except Exception as e:
        return ParseResult(
            attachment_id=attachment_id,
            format="docx",
            raw_text="",
            warnings=[
                ParseWarning(severity="error", message=f"문서 열기 실패: {type(e).__name__}")
            ],
            confidence=0.0,
        )

    blocks: list[ParsedBlock] = []
    full_text: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            blocks.append(ParsedBlock(kind="heading", text=text, level=level))
        else:
            blocks.append(ParsedBlock(kind="paragraph", text=text))
        full_text.append(text)

    # 표는 텍스트만 합쳐서 별도 블록
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            rows.append(row_text)
        if rows:
            table_text = "\n".join(rows)
            blocks.append(ParsedBlock(kind="table", text=table_text))
            full_text.append(table_text)

    return ParseResult(
        attachment_id=attachment_id,
        format="docx",
        raw_text="\n\n".join(full_text),
        blocks=blocks,
        confidence=0.95,
    )
